import PyPDF2
from docx import Document

def pdf_to_word(pdf_path, word_path):
    pdf_file = open(pdf_path, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    
    word_doc = Document()
    
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text = page.extract_text()
        word_doc.add_paragraph(text)
    
    word_doc.save(word_path)
    pdf_file.close()

# Usage example
# pdf_to_word('output_pdf.pdf', 'output.docx')


import os
from docx import Document
from PyPDF2 import PdfWriter
import win32com.client

def convert_doc_to_pdf(doc_file, pdf_file):
    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(doc_file)
    doc.SaveAs(pdf_file, FileFormat=17)
    doc.Close()
    word.Quit()

def convert_docx_to_pdf(docx_file, pdf_file):
    doc = Document(docx_file)
    pdf_writer = PdfWriter()

    for page in doc.inline_shapes:
        if page.type == 3:  # Inline shape type for pictures
            image_stream = page._inline.graphic.graphicData.pic.blipFill.blip.get_attribute("r:embed")
            image_stream = doc.part.related_parts[image_stream].target_part.stream
            pdf_writer.add_page(image_stream)

    for para in doc.paragraphs:
        pdf_writer.add_text(para.text)

    with open(pdf_file, 'wb') as out:
        pdf_writer.write(out)

def convert_to_pdf(input_file, output_file):
    _, file_ext = os.path.splitext(input_file)
    if file_ext.lower() == '.docx':
        convert_docx_to_pdf(input_file, output_file)
    elif file_ext.lower() == '.doc':
        convert_doc_to_pdf(input_file, output_file)
    else:
        print("Unsupported file format.")

# Example usage:
convert_to_pdf('my_doc.doc', 'output.pdf')
# convert_to_pdf('input.doc', 'output.pdf')

